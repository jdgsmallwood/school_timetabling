import os
import pandas
from docx import Document
from pandas import ExcelFile
from pulp import LpProblem, LpMinimize, lpSum, LpVariable, LpStatus, LpInteger, LpBinary
import datetime
import time
from timetabler import app, db, executor
from timetabler.models import *
import timetabler.models
from timetabler.forms import AddTimetableForm

#TIMETABLE CODE
def runtimetable_with_rooms_two_step(STUDENTS, SUBJECTS, TIMES, day, DAYS, TEACHERS, SUBJECTMAPPING, REPEATS,
                                     TEACHERMAPPING,
                                     TUTORAVAILABILITY, maxclasssize, minclasssize, ROOMS, PROJECTORS, PROJECTORROOMS, numroomsprojector, NONPREFERREDTIMES, CAPACITIES):
    '''
    Run the timetabling process and input into the database.

    This process calls the CBCSolver using the PuLP package and then adds the classes to the database.


    :param STUDENTS: should be an array of student names
    :param SUBJECTS: should be an array of subject codes
    :param TIMES: an array of strings representing possible timeslots
    :param day:
    :param DAYS: the days corresponding to the timeslots above
    :param TEACHERS: an array of the names of the tutors
    :param SUBJECTMAPPING: This is a dictionary representing the subjects
                            each tutor is taking
    :param REPEATS: A dictionary of how many repeats each subject has
    :param TEACHERMAPPING: A dictionary of what subject each tutor teachers
    :param TUTORAVAILABILITY:
    :param maxclasssize: An integer representing the maximum class size
    :param minclasssize: An integer representing the minimum class size
    :param nrooms: An integer representing the max allowable concurrent classes
    :param CAPACITIES: A dictionary indexed by room name with the amount of people that each room can contain
    :return: A string representing model status.
    '''
    print("Running solver")
    model = LpProblem('Timetabling', LpMinimize)
    # Create Variables
    print("Creating Variables")
    app.logger.info('Assignment Variables')
    assign_vars = LpVariable.dicts("StudentVariables",
                                   [(i, j, k, m) for m in TEACHERS for j in TEACHERMAPPING[m] for i in SUBJECTMAPPING[j]
                                    for k in TIMES], 0, 1, LpBinary)
    app.logger.info('Subject Variables')
    subject_vars = LpVariable.dicts("SubjectVariables",
                                    [(j, k, m) for m in TEACHERS for j in TEACHERMAPPING[m] for k in TIMES], 0, 1,
                                    LpBinary)

    # c
    app.logger.info('9:30 classes')
    num930classes = LpVariable.dicts("930Classes", [(i) for i in TIMES], lowBound=0, cat=LpInteger)
    # w
    app.logger.info('Days for teachers')
    daysforteachers = LpVariable.dicts("numdaysforteachers", [(i, j) for i in TEACHERS for j in range(len(DAYS))], 0, 1,
                                       LpBinary)
    # p
    daysforteacherssum = LpVariable.dicts("numdaysforteacherssum", [(i) for i in TEACHERS], 0, cat=LpInteger)
    # variables for student clashes
    studenttime = LpVariable.dicts("StudentTime", [(i, j) for i in STUDENTS for j in TIMES], lowBound=0, upBound=1,
                                   cat=LpBinary)
    studentsum = LpVariable.dicts("StudentSum", [(i) for i in STUDENTS], 0, cat=LpInteger)


    projectortime = LpVariable.dicts("ProjectorSum", [(k) for k in TIMES], cat=LpInteger)
    projectorpositive = LpVariable.dicts("ProjectorPositivePart", [(k) for k in TIMES], 0, cat=LpInteger)
    # Count the days that a teacher is rostered on. Make it bigger than a small number times the sum
    # for that particular day.
    for m in TEACHERS:
        app.logger.info('Counting Teachers for ' + m)
        for d in range(len(day)):
            model += daysforteachers[(m, d)] >= 0.1 * lpSum(
                subject_vars[(j, k, m)] for j in TEACHERMAPPING[m] for k in DAYS[day[d]])
            model += daysforteachers[(m, d)] <= lpSum(
                subject_vars[(j, k, m)] for j in TEACHERMAPPING[m] for k in DAYS[day[d]])
    for m in TEACHERS:
        model += daysforteacherssum[(m)] == lpSum(daysforteachers[(m, d)] for d in range(len(day)))

    print("Constraining tutor availability")
    # This bit of code puts in the constraints for the tutor availability.
    # It reads in the 0-1 matrix of tutor availability and constrains that no classes
    # can be scheduled when a tutor is not available.
    # The last column of the availabilities is the tutor identifying number, hence why we have
    # used a somewhat convoluted idea down here.
    for m in TEACHERS:
        for k in TIMES:
            if k not in TUTORAVAILABILITY[m]:
                model += lpSum(subject_vars[(j, k, m)] for j in TEACHERMAPPING[m]) == 0

    # Constraints on subjects for each students
    print("Constraining student subjects")
    for m in TEACHERS:
        for j in TEACHERMAPPING[m]:
            for i in SUBJECTMAPPING[j]:
                model += lpSum(assign_vars[(i, j, k, m)] for k in TIMES) == 1

    # This code means that students cannot attend a tute when a tute is not running
    # But can not attend a tute if they attend a repeat.
    for m in TEACHERS:
        for j in TEACHERMAPPING[m]:
            for i in SUBJECTMAPPING[j]:
                for k in TIMES:
                    model += assign_vars[(i, j, k, m)] <= subject_vars[(j, k, m)]

    # Constraints on which tutor can take each class
    # This goes through each list and either constrains it to 1 or 0 depending if
    # the teacher needs to teach that particular class.
    print("Constraining tutor classes")
    for m in TEACHERS:
        for j in TEACHERMAPPING[m]:
            model += lpSum(subject_vars[(j, k, m)] for k in TIMES) == REPEATS[j]

    # General Constraints on Rooms etc.
    print("Constraining times")
    # For each time cannot exceed number of rooms
    for k in TIMES:
        model += lpSum(subject_vars[(j, k, m)] for m in TEACHERS for j in TEACHERMAPPING[m]) <= len(ROOMS)


    #Number of rooms that need projectors less than the number that have projectors. Want to make this a soft constraint so
    #that it will not affect the feasibility of the model. We'll have a large penalty for exceeding the number of rooms with
    #projectors.
    for k in TIMES:
        model += projectortime[(k)]==(lpSum(subject_vars[(j,k,m)] for m in TEACHERS for j in TEACHERMAPPING[m] if j in PROJECTORS)-numroomsprojector)
        model += projectorpositive[(k)] >= projectortime[(k)]
        model += projectorpositive[(k)] >= 0

    # Teachers can only teach one class at a time
    for k in TIMES:
        for m in TEACHERS:
            model += lpSum(subject_vars[(j, k, m)] for j in TEACHERMAPPING[m]) <= 1
    print("Constraint: Minimize student clashes")
    # STUDENT CLASHES
    for i in STUDENTS:
        for k in TIMES:
            model += studenttime[(i, k)] <= lpSum(
                assign_vars[(i, j, k, m)] for m in TEACHERS for j in TEACHERMAPPING[m] if i in SUBJECTMAPPING[j]) / 2
            model += studenttime[(i, k)] >= 0.3 * (0.5 * lpSum(
                assign_vars[(i, j, k, m)] for m in TEACHERS for j in TEACHERMAPPING[m] if i in SUBJECTMAPPING[j]) - 0.5)
    for i in STUDENTS:
        model += studentsum[(i)] == lpSum(studenttime[(i, k)] for k in TIMES)

    # This minimizes the number of 9:30 classes.
    for i in TIMES:
        if i in NONPREFERREDTIMES:
            model += num930classes[(i)] == lpSum(subject_vars[(j, i, m)] for m in TEACHERS for j in TEACHERMAPPING[m])

        else:
            model += num930classes[(i)] == 0

    print("Setting objective function")

    # Class size constraint
    for m in TEACHERS:
        for j in TEACHERMAPPING[m]:
            for k in TIMES:
                model += lpSum(assign_vars[(i, j, k, m)] for i in SUBJECTMAPPING[j]) >= minclasssize * subject_vars[
                    (j, k, m)]
                model += lpSum(assign_vars[(i, j, k, m)] for i in SUBJECTMAPPING[j]) <= maxclasssize

    # Solving the model
    model += (100 * lpSum(studentsum[(i)] for i in STUDENTS) + lpSum(num930classes[(i)] for i in TIMES) + 500 * lpSum(
        daysforteacherssum[(m)] for m in TEACHERS)+5000*lpSum(projectorpositive[(k)] for k in TIMES))
    print("Solving Model")
    model.solve()
    print("Status:", LpStatus[model.status])
    print("Completed Timetable")

    classpop = {}
    for m in TEACHERS:
        for j in TEACHERMAPPING[m]:
            for k in TIMES:
                if subject_vars[(j,k,m)].varValue == 1:
                    classpop[(j,k,m)] = sum(assign_vars[(i,j,k,m)].varValue for i in SUBJECTMAPPING[j])

    if LpStatus[model.status] == "Optimal":
        print("Allocating Rooms")
        model2 = LpProblem('RoomAllocation', LpMinimize)
        print("Defining Variables")
        subject_vars_rooms = LpVariable.dicts("SubjectVariablesRooms",
                                              [(j, k, m, n) for m in TEACHERS for j in TEACHERMAPPING[m] for k in TIMES for
                                               n in ROOMS if (j,k,m) in classpop.keys()], 0, 1, LpBinary)

        teacher_number_rooms = LpVariable.dicts("NumberRoomsTeacher", [(m, n) for m in TEACHERS for n in ROOMS], 0, 1,
                                                LpBinary)
        teacher_number_rooms_sum = LpVariable.dicts("NumberRoomsTeacherSum", [(m) for m in TEACHERS], 0)

        projector_rooms_sum = LpVariable.dicts("ProjectorRooms", [(j) for j in PROJECTORS])

        populationovershoot = LpVariable.dicts("PopulationOvershoot", [(k,n) for k in TIMES for n in ROOMS])

        poppositive = LpVariable.dicts("PopulationPositivePart", [(k, n) for k in TIMES for n in ROOMS])

        print("Minimizing number of rooms for each tutor")
        for m in TEACHERS:
            for n in ROOMS:
                model2 += teacher_number_rooms[(m, n)] >= 0.01 * lpSum(
                    subject_vars_rooms[(j, k, m, n)] for j in TEACHERMAPPING[m] for k in TIMES if (j,k,m) in classpop.keys())
                model2 += teacher_number_rooms[(m, n)] <= lpSum(
                    subject_vars_rooms[(j, k, m, n)] for j in TEACHERMAPPING[m] for k in TIMES if (j,k,m) in classpop.keys())
        for m in TEACHERS:
            model2 += teacher_number_rooms_sum[(m)] == lpSum(teacher_number_rooms[(m, n)] for n in ROOMS)

        # Rooms must be allocated at times when the classes are running
        print("Constraining Times")
        for m in TEACHERS:
            for j in TEACHERMAPPING[m]:
                for k in TIMES:
                    if (j,k,m) in classpop.keys():
                        model2 += lpSum(subject_vars_rooms[(j, k, m, n)] for n in ROOMS) == subject_vars[(j, k, m)].varValue



        for m in TEACHERS:
            for j in TEACHERMAPPING[m]:
                if j in PROJECTORS:
                    model2 += projector_rooms_sum[(j)] == lpSum(subject_vars_rooms[(j,k,m,n)] for n in PROJECTORROOMS for k in TIMES if (j,k,m) in classpop.keys())

        print("Ensuring Uniqueness")
        # Can only have one class in each room at a time.
        for k in TIMES:
            for n in ROOMS:
                model2 += lpSum(subject_vars_rooms[(j, k, m, n)] for m in TEACHERS for j in TEACHERMAPPING[m] if (j,k,m) in classpop.keys()) <= 1


        print("Accomodating Capacities")

        for k in TIMES:

            for n in ROOMS:



                model2 += populationovershoot[(k,n)] == (lpSum(classpop[(j,k,m)]*subject_vars_rooms[(j,k,m,n)] for m in TEACHERS for j in TEACHERMAPPING[m] if (j,k,m) in classpop.keys()) - CAPACITIES[n])
                #print("Second")
                model2 += poppositive[(k,n)] >= populationovershoot[(k,n)]
                model2 += poppositive[(k,n)] >= 0



        print("Setting Objective Function")
        model2 += lpSum(teacher_number_rooms_sum[(m)] for m in TEACHERS) - 50 * lpSum(projector_rooms_sum[(j)] for j in PROJECTORS) +10 * lpSum(poppositive[(k,n)] for k in TIMES for n in ROOMS)
        print("Solve Room Allocation")
        model2.solve()
        print(LpStatus[model2.status])
        if LpStatus[model2.status] == 'Optimal':
            print("Complete")
            print("Adding to Database")
            timetabler.models.add_classes_to_timetable_twostep(TEACHERS, TEACHERMAPPING, SUBJECTMAPPING, TIMES,
                                                               subject_vars_rooms, assign_vars, ROOMS, classpop)
            print("Status:", LpStatus[model2.status])
    return LpStatus[model.status]



def preparetimetable(addtonewtimetable=False):
    '''
    Get timetable data and then execute the timetabling program.

    :param addtonewtimetable: Whether this should be added to a new timetable and set as default.
    :return: The view timetable page.

    TESTED
    '''
    print("Preparing Timetable")

    (STUDENTS, SUBJECTS, TIMES, day, DAYS, TEACHERS, SUBJECTMAPPING, REPEATS, TEACHERMAPPING,
     TUTORAVAILABILITY, maxclasssize, minclasssize, ROOMS,PROJECTORS, PROJECTORROOMS, numroomsprojector, NONPREFERREDTIMES, CAPACITIES) = timetabler.models.get_timetable_data(rooms=True)

    print("Everything ready")
    executor.submit(runtimetable_with_rooms_two_step, STUDENTS, SUBJECTS, TIMES, day, DAYS, TEACHERS, SUBJECTMAPPING,
                    REPEATS, TEACHERMAPPING,
                    TUTORAVAILABILITY, maxclasssize, minclasssize, ROOMS,PROJECTORS, PROJECTORROOMS, numroomsprojector, NONPREFERREDTIMES,CAPACITIES)


    form = AddTimetableForm()
    return render_template("viewtimetable.html", form=form)




def allowed_file(filename):
    '''
    Checks whether the uploaded file has an allowed extension.
    :param filename: The filename to check
    :return: True/False

    TESTED
    '''
    return '.' in filename and \
           filename.rsplit('.', 1)[1] in app.config['ALLOWED_EXTENSIONS']


def upload(file):
    '''
    Save the uploaded file to the UPLOAD_FOLDER directory.

    :param file: The file to upload
    :return: Filename of uploaded file in upload folder
    '''
    if file and allowed_file(file.filename):
        # Make the filename safe, remove unsupported chars
        filename = file.filename
        # Ensure unique filename
        unique_filename = time.strftime("%Y-%m-%d_%H%M%S") + filename
        # Move the file form the temporal folder to
        # the upload folder we setup
        path_to_file = os.path.join(app.config['UPLOAD_FOLDER'], unique_filename)
        file.save(path_to_file)
        return path_to_file


def checkboxvalue(checkbox):
    '''
    Get value of checkbox.

    :param checkbox: Input from request.form
    :return: 1 if ticked, 0 if not.

    TESTED
    '''
    if (checkbox != None):
        return 1
    else:
        return 0


def read_excel(filename):
    '''
    Read Excel File provided by filename.

    :param filename - path to an Excel file:
    :return: pandas dataframe
    '''
    xl = ExcelFile(filename)
    df = xl.parse(xl.sheet_names[0])
    return df

def read_csv(filename):
    '''
    Read CSV File provided by filename.

    :param filename - path to an CSV file:
    :return: pandas dataframe
    '''
    return pandas.read_csv(filename)


def create_roll(students, subject, timeslot, room):
    path_to_file = app.config['UPLOAD_FOLDER'] + '/roll_' + \
        subject.subcode + '_' + time.strftime("%Y-%m-%d_%H%M%S") + '.docx'

    document = Document()

    document.add_heading(subject.subname, 0)

    document.add_paragraph('Timeslot: ' + timeslot.day + " " + timeslot.time)
    if room is not None:
        document.add_paragraph('Room: ' + room.name)

    table = document.add_table(rows=1, cols=12)
    table.style = 'TableGrid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Name'
    hdr_cells[1].text = '1'
    hdr_cells[2].text = '2'
    hdr_cells[3].text = '3'
    hdr_cells[4].text = '4'
    hdr_cells[5].text = '5'
    hdr_cells[6].text = '6'
    hdr_cells[7].text = '7'
    hdr_cells[8].text = '8'
    hdr_cells[9].text = '9'
    hdr_cells[10].text = '10'
    hdr_cells[11].text = '11'
    for item in students:
        row_cells = table.add_row().cells
        row_cells[0].text = str(item.name)
    document.save(path_to_file)
    return path_to_file


def create_excel(data):
    path_to_file = app.config['UPLOAD_FOLDER'] + '/timetable' + time.strftime("%Y-%m-%d_%H%M%S") + '.xlsx'
    writer = pandas.ExcelWriter(path_to_file, engine='xlsxwriter')
    data.to_excel(writer, sheet_name='Timetable', index=False)
    writer.save()
    return path_to_file


def format_timetable_data_for_export():
    timeslots = timetabler.models.get_all_timeslots()
    timeslots = sorted(timeslots, key=attrgetter('daynumeric', 'time'))
    timetable = []
    for i in range(len(timeslots)):
        timeslot = timeslots[i]
        classes = timeslot.timetabledclasses
        for timeclass in classes:
            if timeclass.room is not None:
                room = timeclass.room.name
            else:
                room = ""
            if timeclass.tutor is not None:
                tutor = timeclass.tutor.name
            else:
                tutor = ""

            timetable.append((
                            timeclass.timeslot.day + ' ' + timeclass.timeslot.time,
                            timeclass.subject.subcode,
                            timeclass.subject.subname,
                            tutor,
                            room))

    timetable = pandas.DataFrame(timetable)
    timetable.columns = ['Time', 'SubjectCode', 'SubjectName', 'Tutor', 'Room']
    return timetable


def format_student_timetable_data_for_export():
    students = timetabler.models.Student.get_all()
    timetable = []
    for student in students:
        for timeclass in student.timetabledclasses:
            if timeclass.room is not None:
                room = timeclass.room.name
            else:
                room = ""
            if timeclass.tutor is not None:
                tutor = timeclass.tutor.name
            else:
                tutor = ""
            timetable.append((
                student.studentcode,
                student.name,
                timeclass.subject.subcode,
                timeclass.subject.subname,
                timeclass.timeslot.day + ' ' + timeclass.timeslot.time,
                tutor,
                room
            ))

    timetable = pandas.DataFrame(timetable)
    timetable.columns = ['StudentId', 'StudentName', 'SubjectCode', 'SubjectName', 'Time', 'Tutor', 'Room']
    return timetable


def format_tutor_hours_for_export(hours):
    hours = list(hours)
    hours = pandas.DataFrame(hours)
    hours.columns = ['Name', 'Initial Tutorials', 'Repeat Tutorials']
    return hours


def convert_to_datetime(value):
    return datetime.datetime.strptime(value, '%Y-%m-%dT%H:%M')


def convert_datetime_to_string(value):
    return value.strftime('%d-%m-%Y %H:%M')
